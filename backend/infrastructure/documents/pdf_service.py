from __future__ import annotations

import time
import zipfile
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutureTimeoutError
from pathlib import Path

import docx
import pypdf

from infrastructure.documents.models import (
    DocumentExtractionResult,
    DocumentExtractionStatus,
    DocumentPage,
    DocumentProcessingError,
    DocumentTable,
)
from infrastructure.documents.ocr import OCREngine, RapidOCREngine
from infrastructure.documents.quality import TextQuality, classify_text_quality, normalize_extracted_text


PDF_MIME_TYPES = {"application/pdf"}
DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/zip",
}
MAX_DOCX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
MAX_OCR_PAGES = 20


class DocumentService:
    def __init__(self, *, ocr_engine: OCREngine | None = None, ocr_timeout_seconds: float = 30.0) -> None:
        self._ocr_engine = ocr_engine or RapidOCREngine()
        self.ocr_timeout_seconds = ocr_timeout_seconds

    def extract_document(
        self,
        file_path: str,
        filename: str | None = None,
        *,
        content_type: str | None = None,
    ) -> DocumentExtractionResult:
        name = filename or Path(file_path).name
        source_type = self._validate_document(file_path, name, content_type)
        return self._extract_pdf(file_path) if source_type == "pdf" else self._extract_docx(file_path)

    def extract_text(self, file_path: str, filename: str | None = None) -> str:
        return self.extract_document(file_path, filename).text

    def extract_pdf_text(self, file_path: str) -> str:
        return self._extract_pdf(file_path).text

    def extract_docx_text(self, file_path: str) -> str:
        return self._extract_docx(file_path).text

    extract_pdf_text_direct = extract_pdf_text
    extract_docx_text_direct = extract_docx_text

    def _validate_document(self, file_path: str, filename: str, content_type: str | None) -> str:
        suffix = Path(filename).suffix.casefold()
        if suffix not in {".pdf", ".docx"}:
            raise DocumentProcessingError(
                "unsupported_file_type", "Only PDF and DOCX resumes are supported.", status_code=415
            )
        with open(file_path, "rb") as stream:
            header = stream.read(1024)
        detected = "pdf" if b"%PDF-" in header else "docx" if header.startswith(b"PK") else None
        expected = suffix[1:]
        if detected != expected:
            raise DocumentProcessingError(
                "file_type_mismatch",
                "The file content does not match its PDF or DOCX filename.",
                status_code=415,
            )
        normalized_mime = (content_type or "").split(";", 1)[0].strip().casefold()
        allowed_mimes = PDF_MIME_TYPES if expected == "pdf" else DOCX_MIME_TYPES
        if normalized_mime and normalized_mime not in allowed_mimes:
            raise DocumentProcessingError(
                "file_type_mismatch", "The declared file type does not match the document.", status_code=415
            )
        if expected == "docx":
            try:
                with zipfile.ZipFile(file_path) as archive:
                    members = set(archive.namelist())
                    expanded_size = sum(item.file_size for item in archive.infolist())
                if not {"[Content_Types].xml", "word/document.xml"}.issubset(members):
                    raise ValueError("missing Word package members")
                if expanded_size > MAX_DOCX_UNCOMPRESSED_BYTES:
                    raise ValueError("expanded Word package is too large")
            except (OSError, zipfile.BadZipFile, ValueError) as error:
                raise DocumentProcessingError(
                    "invalid_document", "The DOCX file is malformed or cannot be opened.", status_code=422
                ) from error
        return expected

    def _extract_pdf(self, file_path: str) -> DocumentExtractionResult:
        """Extract text from a PDF.

        Primary path: ``pymupdf4llm.to_markdown()`` — preserves layout and
        section structure, which improves LLM extraction quality.

        Fallback path: ``pypdf`` native extraction + OCR — used when
        pymupdf4llm is not installed or returns less than 50 characters
        (e.g. scanned/image-only PDFs that pymupdf can't OCR by itself).
        """
        # ----------------------------------------------------------------
        # Primary path: pymupdf4llm layout-aware markdown extraction
        # ----------------------------------------------------------------
        md_text = self._try_pymupdf4llm(file_path)
        if md_text:
            return DocumentExtractionResult(
                text=md_text,
                source_type="pdf",
                character_count=len(md_text),
                extraction_method="pymupdf4llm",
                status=DocumentExtractionStatus.COMPLETE,
            )

        # ----------------------------------------------------------------
        # Fallback path: pypdf + OCR
        # ----------------------------------------------------------------
        warnings: list[str] = []
        try:
            with open(file_path, "rb") as stream:
                reader = pypdf.PdfReader(stream)
                if reader.is_encrypted and not reader.decrypt(""):
                    raise DocumentProcessingError(
                        "encrypted_document", "Password-protected PDF files are not supported.", status_code=422
                    )
                raw_pages: list[str] = []
                for page in reader.pages:
                    try:
                        raw_pages.append(page.extract_text() or "")
                    except Exception:
                        raw_pages.append("")
                        warnings.append("page_parse_failed")
        except DocumentProcessingError:
            raise
        except Exception as error:
            raise DocumentProcessingError(
                "invalid_document", "The PDF file is malformed or cannot be opened.", status_code=422
            ) from error

        page_methods = ["native_pdf" for _ in raw_pages]
        page_warnings: list[list[str]] = [[] for _ in raw_pages]
        if classify_text_quality(raw_pages) in {TextQuality.IMAGE_ONLY, TextQuality.SPARSE}:
            engine = self._ocr_engine
            deadline = time.monotonic() + self.ocr_timeout_seconds
            for index, native_text in enumerate(raw_pages):
                if sum(character.isalnum() for character in native_text) >= 50:
                    continue
                if index >= MAX_OCR_PAGES:
                    page_warnings[index].append("ocr_page_limit_reached")
                    warnings.append("ocr_page_limit_reached")
                    continue
                try:
                    remaining = deadline - time.monotonic()
                    if remaining <= 0:
                        raise TimeoutError("OCR document timeout")
                    ocr_text = self._recognize_with_timeout(
                        engine,
                        self._render_pdf_page(file_path, index),
                        timeout_seconds=remaining,
                    )
                    if ocr_text.strip():
                        raw_pages[index] = ocr_text
                        page_methods[index] = "ocr"
                    else:
                        page_warnings[index].append("ocr_empty")
                        warnings.append("ocr_empty")
                except TimeoutError:
                    page_warnings[index].append("ocr_timeout")
                    warnings.append("ocr_timeout")
                except Exception:
                    page_warnings[index].append("ocr_failed")
                    warnings.append("ocr_failed")

        normalized_pages = [normalize_extracted_text(text) for text in raw_pages]
        text = normalize_extracted_text("\n".join(normalized_pages))
        if len(text) < 50:
            raise DocumentProcessingError(
                "no_extractable_text",
                "Could not recover enough readable text from the document.",
                status_code=422,
                warnings=list(dict.fromkeys(warnings)),
            )
        methods = set(page_methods)
        extraction_method = "mixed" if len(methods) > 1 else next(iter(methods), "native_pdf")
        unique_warnings = list(dict.fromkeys(warnings))
        is_partial = bool(unique_warnings)
        return DocumentExtractionResult(
            text=text,
            source_type="pdf",
            page_count=len(raw_pages),
            character_count=len(text),
            extraction_method=extraction_method,
            status=DocumentExtractionStatus.PARTIAL if is_partial else DocumentExtractionStatus.COMPLETE,
            is_partial=is_partial,
            warnings=unique_warnings,
            pages=[
                DocumentPage(
                    page_number=index + 1,
                    text=normalized_pages[index],
                    extraction_method=page_methods[index],
                    warnings=page_warnings[index],
                )
                for index in range(len(raw_pages))
            ],
        )

    @staticmethod
    def _try_pymupdf4llm(file_path: str) -> str:
        """Attempt layout-aware markdown extraction via pymupdf4llm.

        Returns the markdown string on success, or an empty string if the
        library is unavailable or the result is too short to be useful.
        """
        try:
            import pymupdf4llm  # optional dependency

            md = pymupdf4llm.to_markdown(file_path)
            cleaned = md.strip()
            if len(cleaned) >= 50:
                return cleaned
            return ""
        except Exception:
            return ""

    def _extract_docx(self, file_path: str) -> DocumentExtractionResult:
        try:
            document = docx.Document(file_path)
        except Exception as error:
            raise DocumentProcessingError(
                "invalid_document", "The DOCX file is malformed or cannot be opened.", status_code=422
            ) from error
        text_parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        tables: list[DocumentTable] = []
        for table in document.tables:
            rows: list[list[str]] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(cells)
                    text_parts.append(" | ".join(cell for cell in cells if cell))
            if rows:
                tables.append(DocumentTable(rows=rows))
        text = normalize_extracted_text("\n".join(text_parts))
        if len(text) < 50:
            raise DocumentProcessingError(
                "no_extractable_text", "Could not recover enough readable text from the document.", status_code=422
            )
        return DocumentExtractionResult(
            text=text,
            source_type="docx",
            character_count=len(text),
            extraction_method="docx",
            status=DocumentExtractionStatus.COMPLETE,
            tables=tables,
        )

    def _render_pdf_page(self, file_path: str, page_index: int) -> bytes:
        import pymupdf

        with pymupdf.open(file_path) as document:
            pixmap = document[page_index].get_pixmap(matrix=pymupdf.Matrix(2, 2), alpha=False)
            return pixmap.tobytes("png")

    def _recognize_with_timeout(
        self, engine: OCREngine, image_bytes: bytes, *, timeout_seconds: float
    ) -> str:
        executor = ThreadPoolExecutor(max_workers=1)
        future = executor.submit(engine.recognize, image_bytes)
        try:
            return future.result(timeout=timeout_seconds)
        except FutureTimeoutError as error:
            future.cancel()
            raise TimeoutError("OCR timeout") from error
        finally:
            executor.shutdown(wait=False, cancel_futures=True)

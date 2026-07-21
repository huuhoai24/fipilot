import calendar
import json
import math
import os
import re
from datetime import date
from difflib import SequenceMatcher

import docx
import httpx
import pypdf
from openai import AsyncOpenAI


CURRENT_YEAR = date.today().year


SKILL_TAXONOMY = {
    "Python": ["python", "py"],
    "Java": ["java", "spring boot", "spring"],
    "JavaScript": ["javascript", "js", "ecmascript"],
    "TypeScript": ["typescript", "ts"],
    "React": ["react", "reactjs", "react.js", "next.js", "nextjs"],
    "Vue": ["vue", "vue.js", "nuxt"],
    "Angular": ["angular"],
    "Node.js": ["node.js", "nodejs", "node", "express", "nestjs"],
    "C#": ["c#", ".net", "asp.net"],
    "Go": ["golang", "go developer"],
    "PHP": ["php", "laravel"],
    "SQL": ["sql", "mysql", "postgresql", "postgres", "mssql", "oracle"],
    "NoSQL": ["mongodb", "mongo", "redis", "cassandra", "dynamodb"],
    "API": ["api", "rest", "restful", "graphql", "grpc"],
    "Docker": ["docker", "container"],
    "Kubernetes": ["kubernetes", "k8s"],
    "AWS": ["aws", "amazon web services"],
    "Azure": ["azure"],
    "GCP": ["gcp", "google cloud"],
    "Linux": ["linux", "ubuntu", "bash", "shell"],
    "Git": ["git"],
    "GitHub": ["github"],
    "GitLab": ["gitlab"],
    "Bitbucket": ["bitbucket"],
    "CI/CD": ["ci/cd", "cicd", "jenkins", "github actions", "gitlab ci"],
    "Testing": ["testing", "unit test", "integration test", "pytest", "jest", "selenium", "cypress"],
    "Playwright": ["playwright"],
    "QA Automation": ["qa automation", "automation test", "automated testing", "selenium", "cypress"],
    "HTML": ["html", "html5"],
    "CSS": ["css", "css3", "sass", "scss", "tailwind", "bootstrap"],
    "AI/ML": ["ai/ml", "artificial intelligence", "ai powered", "ai-powered", "ai model"],
    "Machine Learning": ["machine learning", "scikit-learn", "sklearn"],
    "Deep Learning": ["deep learning", "neural network", "tensorflow", "keras"],
    "PyTorch": ["pytorch", "torch"],
    "CNN": ["cnn", "convolutional neural network"],
    "Transformer": ["transformer", "transformers", "bert", "gpt"],
    "OCR": ["ocr", "optical character recognition"],
    "LLM": ["llm", "large language model", "openai", "ollama", "llamaindex"],
    "LLM APIs": ["llm api", "llm apis", "openai api", "gemini api", "ollama api"],
    "LangChain": ["langchain"],
    "LangGraph": ["langgraph"],
    "RAG": ["rag", "retrieval augmented generation", "vector database", "embedding", "faiss", "qdrant", "pinecone"],
    "NLP": ["nlp", "natural language processing", "spacy", "nltk"],
    "FastAPI": ["fastapi", "fast api"],
    "Computer Vision": ["computer vision", "image processing", "object detection"],
    "OpenCV": ["opencv", "open cv"],
    "Ultralytics": ["ultralytics", "yolo", "yolov8", "yolov10", "yolov11"],
    "CLIP": ["clip", "openai clip"],
    "Pandas": ["pandas"],
    "Matplotlib": ["matplotlib", "pyplot"],
    "MLOps": ["mlops", "model serving", "mlflow", "kubeflow", "airflow"],
    "Data Engineering": ["data engineering", "etl", "elt", "data pipeline", "data warehouse"],
    "Spark": ["spark", "pyspark", "databricks"],
    "Airflow": ["airflow", "dag"],
    "Kafka": ["kafka", "streaming"],
    "Power BI": ["power bi", "powerbi"],
    "Tableau": ["tableau"],
    "Business Analysis": ["business analyst", "business analysis", "brd", "user story", "use case", "bpmn"],
    "Security": ["security", "cybersecurity", "penetration testing", "pentest", "owasp", "siem"],
}


ROLE_PROFILES = {
    "Backend Developer": {
        "title": ["backend", "back-end", "server side", "api developer"],
        "skills": ["Python", "Java", "Node.js", "C#", "Go", "PHP", "SQL", "NoSQL", "API", "Docker", "Git", "GitHub"],
    },
    "Frontend Developer": {
        "title": ["frontend", "front-end", "web ui", "ui developer"],
        "skills": ["JavaScript", "TypeScript", "React", "Vue", "Angular", "HTML", "CSS", "Testing", "Playwright", "Git", "GitHub"],
    },
    "Fullstack Developer": {
        "title": ["fullstack", "full-stack", "full stack"],
        "skills": ["JavaScript", "TypeScript", "React", "Node.js", "Python", "Java", "SQL", "API", "Docker"],
    },
    "Web Developer": {
        "title": ["web developer", "website developer"],
        "skills": ["JavaScript", "TypeScript", "React", "Vue", "HTML", "CSS", "Node.js", "PHP", "SQL"],
    },
    "Software Engineer": {
        "title": ["software engineer", "software developer", "developer", "programmer"],
        "skills": ["Python", "Java", "JavaScript", "TypeScript", "SQL", "API", "Git", "GitHub", "Testing", "Playwright", "Docker"],
    },
    "AI Engineer": {
        "title": [
            "ai engineer",
            "artificial intelligence",
            "machine learning engineer",
            "ml engineer",
            "nlp engineer",
            "computer vision engineer",
            "ai ml",
        ],
        "skills": [
            "Python",
            "AI/ML",
            "Machine Learning",
            "Deep Learning",
            "PyTorch",
            "CNN",
            "Transformer",
            "OCR",
            "LLM",
            "LLM APIs",
            "LangChain",
            "LangGraph",
            "RAG",
            "NLP",
            "FastAPI",
            "Computer Vision",
            "OpenCV",
            "Ultralytics",
            "CLIP",
            "MLOps",
            "SQL",
            "Git",
            "GitHub",
        ],
    },
    "Data Scientist": {
        "title": ["data scientist", "machine learning scientist"],
        "skills": ["Python", "SQL", "Machine Learning", "Deep Learning", "PyTorch", "NLP", "Pandas", "Matplotlib", "Power BI", "Tableau"],
    },
    "Data Engineer": {
        "title": ["data engineer", "etl developer", "big data engineer"],
        "skills": ["Python", "SQL", "Data Engineering", "Spark", "Airflow", "Kafka", "AWS", "GCP", "Azure"],
    },
    "DevOps Engineer": {
        "title": ["devops", "site reliability", "sre", "cloud engineer"],
        "skills": ["Docker", "Kubernetes", "AWS", "Azure", "GCP", "Linux", "CI/CD", "Git"],
    },
    "Cloud Engineer": {
        "title": ["cloud engineer", "cloud architect"],
        "skills": ["AWS", "Azure", "GCP", "Docker", "Kubernetes", "Linux", "CI/CD"],
    },
    "Tester": {
        "title": ["tester", "qa", "quality assurance", "test engineer"],
        "skills": ["Testing", "QA Automation", "Playwright", "SQL", "API", "Selenium", "Git", "GitHub"],
    },
    "Business Analyst": {
        "title": ["business analyst", "product analyst"],
        "skills": ["Business Analysis", "SQL", "Power BI", "Tableau", "API"],
    },
    "Cybersecurity Analyst": {
        "title": ["cybersecurity", "security analyst", "soc analyst"],
        "skills": ["Security", "Linux", "Python", "SQL", "AWS"],
    },
}


SECTION_ALIASES = {
    "summary": ["summary", "professional summary", "profile", "objective", "about me", "career objective"],
    "skills": ["skills", "technical skills", "technologies", "tools", "tech stack", "competencies", "software"],
    "experience": [
        "experience",
        "work experience",
        "working experience",
        "working experience highlights",
        "professional experience",
        "employment",
        "work history",
    ],
    "projects": ["projects", "personal projects", "academic projects", "portfolio", "project and practice"],
    "education": [
        "education",
        "education certificate",
        "education and certificate",
        "education and training",
        "educational background",
        "academic background",
        "training",
        "courses",
        "qualification",
        "qualifications",
    ],
    "certifications": ["certifications", "certificates", "certificate and award", "licenses"],
    "hobbies": ["hobbies", "interests"],
}


TABLE_LABELS = {
    "project name",
    "duration",
    "position",
    "positions",
    "role",
    "title",
    "general information",
    "description",
    "project scope",
    "technology used",
    "technologies used",
    "team size",
    "customer",
    "client",
    "responsibilities",
}


ROLE_TITLE_TERMS = {
    "engineer",
    "developer",
    "analyst",
    "tester",
    "qa",
    "devops",
    "scientist",
    "architect",
    "manager",
    "lead",
    "specialist",
    "consultant",
    "administrator",
}


GENERIC_ROLE_ALIASES = {"developer", "engineer", "analyst", "tester", "programmer"}

EDUCATION_DATE_CONTEXT_TERMS = {
    "academic",
    "bachelor",
    "campus",
    "college",
    "degree",
    "diploma",
    "education",
    "expected graduation",
    "faculty",
    "gpa",
    "graduation",
    "major",
    "school",
    "student",
    "university",
}

WORK_DATE_CONTEXT_TERMS = {
    "company",
    "contract",
    "developer",
    "employment",
    "engineer",
    "experience",
    "freelance",
    "full time",
    "full-time",
    "intern",
    "internship",
    "part time",
    "part-time",
    "position",
    "remote",
    "role",
    "work",
}


MONTHS = {
    name.lower(): index for index, name in enumerate(calendar.month_name) if name
}
MONTHS.update({name.lower(): index for index, name in enumerate(calendar.month_abbr) if name})


class CVExtractor:
    def __init__(self):
        self.llm_model = os.environ.get("OLLAMA_CV_MODEL") or os.environ.get("CORE_MODEL", "gemma4:e2b")
        self.llm_provider = os.environ.get("CV_LLM_PROVIDER", "ollama").lower()
        self.remote_model_url = os.environ.get("REMOTE_MODEL_URL", "").rstrip("/")
        self.remote_model_token = os.environ.get("REMOTE_MODEL_TOKEN", "")
        self.llm_client = AsyncOpenAI(
            api_key=os.environ.get("OLLAMA_API_KEY") or os.environ.get("CORE_API_KEY", "ollama"),
            base_url=os.environ.get("OLLAMA_BASE_URL") or os.environ.get("CORE_BASE_URL", "http://localhost:11434/v1"),
            http_client=httpx.AsyncClient(verify=False, timeout=45),
        )

    def extract_text(self, file_path: str, filename: str) -> str:
        text = ""
        ext = filename.split(".")[-1].lower()
        try:
            if ext == "pdf":
                with open(file_path, "rb") as f:
                    reader = pypdf.PdfReader(f)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            elif ext in ["docx", "doc"]:
                document = docx.Document(file_path)
                for para in document.paragraphs:
                    if para.text.strip():
                        text += para.text + "\n"
                for table in document.tables:
                    for row in table.rows:
                        cells = self._dedupe_cells([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if cells:
                            text += " | ".join(cells) + "\n"
        except Exception as e:
            print(f"Error parsing CV: {e}")
        return self._normalize_whitespace(text)

    async def parse_cv(self, text: str) -> dict:
        return self.parse_cv_sync(text)

    async def parse_cv_with_llm(self, text: str) -> dict:
        resume_text = self._normalize_whitespace(text)[:12000]
        prompt = f"""
Extract the candidate profile from this CV and return only one valid JSON object.

Required schema:
{{
  "candidate_name": "string",
  "years_experience": number,
  "skills": ["string"],
  "education": "string",
  "recent_role": "string",
  "inferred_level": integer,
  "role_fit": "string",
  "confidence": number
}}

Rules:
- inferred_level must be 1 for fresher/junior, 2 for 2-4 years, 3 for senior/lead, 4 for principal/architect/manager.
- role_fit should be one common IT role such as Software Engineer, Backend Developer, Frontend Developer, AI Engineer, Data Scientist, Data Engineer, DevOps Engineer, Tester, Business Analyst, Cybersecurity Analyst.
- education must contain only school, university, college, academy, training center, degree, major, course, or expected graduation date. Do not put language proficiency, skills, work experience, or project names in education.
- Keep education under 180 characters.
- Keep recent_role under 180 characters.
- Return at most 20 skills.
- confidence must be between 0 and 1.
- If a field is missing, infer conservatively from the CV text.
- Do not include markdown, comments, or trailing text.

CV:
{resume_text}
"""
        if self.llm_provider in {"remote", "model_server"}:
            if not self.remote_model_url:
                raise RuntimeError("REMOTE_MODEL_URL is not configured for CV_LLM_PROVIDER=remote")
            headers = {}
            if self.remote_model_token:
                headers["Authorization"] = f"Bearer {self.remote_model_token}"
            payload = {
                "model": self.llm_model,
                "messages": [
                    {"role": "system", "content": "You are a precise resume extraction engine. Return JSON only."},
                    {"role": "user", "content": prompt},
                ],
                "json_mode": True,
                "temperature": 0.1,
                "max_new_tokens": 700,
            }
            async with httpx.AsyncClient(timeout=180) as client:
                response = await client.post(f"{self.remote_model_url}/llm", json=payload, headers=headers)
                response.raise_for_status()
                content = response.json().get("text") or "{}"
            profile = await self._load_llm_profile_json(content, prompt, headers=headers)
            return self._normalize_profile(profile, extraction_method=f"remote_llm:{self.llm_model}")

        response = await self.llm_client.chat.completions.create(
            model=self.llm_model,
            messages=[
                {"role": "system", "content": "You are a precise resume extraction engine. Return JSON only."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.1,
            response_format={"type": "json_object"},
        )
        content = response.choices[0].message.content or "{}"
        return self._normalize_profile(self._loads_json_object(content), extraction_method=f"llm:{self.llm_model}")

    async def _load_llm_profile_json(self, content: str, original_prompt: str, headers: dict) -> dict:
        try:
            return self._loads_json_object(content)
        except ValueError:
            repair_prompt = (
                "The previous answer was invalid JSON. Repair it and return only one valid JSON object "
                "matching the candidate profile schema. Do not add markdown.\n\n"
                f"Original extraction task:\n{original_prompt[:6000]}\n\nInvalid JSON/text:\n{content[:4000]}"
            )
            payload = {
                "model": self.llm_model,
                "messages": [
                    {"role": "system", "content": "You repair malformed JSON for resume extraction. Return JSON only."},
                    {"role": "user", "content": repair_prompt},
                ],
                "json_mode": True,
                "temperature": 0.0,
                "max_new_tokens": 700,
            }
            async with httpx.AsyncClient(timeout=180) as client:
                response = await client.post(f"{self.remote_model_url}/llm", json=payload, headers=headers)
                response.raise_for_status()
                repaired = response.json().get("text") or "{}"
            return self._loads_json_object(repaired)

    def _loads_json_object(self, content: str) -> dict:
        text = self._extract_json_object(content)
        text = text.replace("\u201c", '"').replace("\u201d", '"').replace("\u2018", "'").replace("\u2019", "'")
        text = re.sub(r",\s*([}\]])", r"\1", text)
        try:
            value = json.loads(text)
        except json.JSONDecodeError as exc:
            snippet = text[max(0, exc.pos - 120): exc.pos + 120]
            raise ValueError(f"Invalid JSON from CV LLM near char {exc.pos}: {snippet}") from exc
        if not isinstance(value, dict):
            raise ValueError("CV LLM returned JSON, but it was not an object")
        return value

    def _extract_json_object(self, content: str) -> str:
        text = (content or "").strip()
        text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.I)
        text = re.sub(r"\s*```$", "", text)
        start = text.find("{")
        if start < 0:
            raise ValueError("CV LLM did not return a JSON object")
        depth = 0
        in_string = False
        escape = False
        for index in range(start, len(text)):
            char = text[index]
            if in_string:
                if escape:
                    escape = False
                elif char == "\\":
                    escape = True
                elif char == '"':
                    in_string = False
                continue
            if char == '"':
                in_string = True
            elif char == "{":
                depth += 1
            elif char == "}":
                depth -= 1
                if depth == 0:
                    return text[start:index + 1]
        raise ValueError("CV LLM returned an incomplete JSON object")

    def parse_cv_sync(self, text: str) -> dict:
        normalized_text = self._normalize_whitespace(text)
        sections = self._split_sections(normalized_text)
        skills = self._extract_skills(normalized_text, sections)
        recent_role = self._extract_recent_role(normalized_text, sections)
        years_experience = self._estimate_years_experience(normalized_text, sections)
        role_fit, role_confidence = self._infer_role(normalized_text, recent_role, skills)
        role_fit, role_confidence = self._refine_role_fit(normalized_text, sections, recent_role, skills, role_fit, role_confidence)
        inferred_level = self._infer_level(years_experience, recent_role, normalized_text)

        confidence_parts = [
            0.2 if self._extract_candidate_name(normalized_text) != "Candidate" else 0.0,
            min(len(skills) / 8, 1.0) * 0.3,
            min(years_experience / 3, 1.0) * 0.2 if years_experience else 0.05,
            role_confidence * 0.3,
        ]

        return self._normalize_profile({
            "candidate_name": self._extract_candidate_name(normalized_text),
            "years_experience": years_experience,
            "skills": skills or ["Not Found"],
            "education": self._extract_education(sections),
            "recent_role": recent_role,
            "inferred_level": inferred_level,
            "role_fit": role_fit,
            "confidence": round(min(sum(confidence_parts), 0.98), 2),
            "extraction_method": "rule_based",
        })

    def _refine_role_fit(
        self,
        text: str,
        sections: dict,
        recent_role: str,
        skills: list,
        role_fit: str,
        confidence: float,
    ) -> tuple[str, float]:
        normalized = self._clean_heading(
            "\n".join([
                text[:5000],
                sections.get("education", ""),
                sections.get("projects", ""),
                recent_role,
            ])
        )
        skill_set = set(skills)
        ai_skill_hits = len(skill_set & {
            "AI/ML",
            "Machine Learning",
            "Deep Learning",
            "PyTorch",
            "CNN",
            "Transformer",
            "OCR",
            "LLM",
            "LLM APIs",
            "LangChain",
            "LangGraph",
            "RAG",
            "NLP",
            "Computer Vision",
            "OpenCV",
            "Ultralytics",
            "CLIP",
        })
        has_ai_degree = any(
            phrase in normalized
            for phrase in [
                "bachelor of artificial intelligence",
                "artificial intelligence",
                "ai campus",
                "major ai",
                "major artificial intelligence",
            ]
        )
        has_ai_project = any(
            phrase in normalized
            for phrase in [
                "anti bot ai crawler",
                "ai crawler",
                "computer vision",
                "ocr",
                "rag",
                "pytorch",
                "ultralytics",
                "opencv",
                "langchain",
                "langgraph",
            ]
        )

        if has_ai_degree and ai_skill_hits >= 2:
            return "AI Engineer", max(confidence, 0.85)
        if ai_skill_hits >= 5 and has_ai_project:
            return "AI Engineer", max(confidence, 0.82)
        if ai_skill_hits >= 7:
            return "AI Engineer", max(confidence, 0.8)
        return role_fit, confidence

    def _normalize_profile(self, profile: dict, extraction_method: str | None = None) -> dict:
        skills = profile.get("skills") if isinstance(profile.get("skills"), list) else []
        clean_skills = [str(skill).strip() for skill in skills if str(skill).strip()]
        years = self._coerce_float(profile.get("years_experience"), 0.0)
        level = int(self._coerce_float(profile.get("inferred_level"), 1))
        confidence = self._coerce_float(profile.get("confidence"), 0.7)

        return {
            "candidate_name": str(profile.get("candidate_name") or "Candidate").strip()[:120],
            "years_experience": round(max(years, 0.0), 1),
            "skills": clean_skills[:32] or ["Not Found"],
            "education": self._sanitize_education_summary(profile.get("education")),
            "recent_role": str(profile.get("recent_role") or "Not Found").strip()[:140],
            "inferred_level": max(1, min(level, 4)),
            "role_fit": str(profile.get("role_fit") or "Software Engineer").strip()[:80],
            "confidence": round(max(0.0, min(confidence, 0.99)), 2),
            "extraction_method": extraction_method or profile.get("extraction_method") or "rule_based",
        }

    def _sanitize_education_summary(self, value) -> str:
        text = str(value or "").strip()
        if not text or text.lower() in {"not found", "none", "n/a", "null"}:
            return "Not Found"

        graduation = self._extract_expected_graduation(text)
        text = re.sub(
            r"\b(?:english|japanese|chinese|korean|french|german|spanish|vietnamese)\s*:?\s*"
            r"(?:basic|beginner|intermediate|advanced|native|fluent|professional|business|n[1-5]|[a-c][1-2])"
            r"(?:\s+level)?\b",
            "",
            text,
            flags=re.I,
        )
        text = re.sub(r"\b(?:ielts|toeic|toefl|jlpt|gpa)\s*:?\s*[a-z0-9./+-]+\b", "", text, flags=re.I)
        text = re.sub(r"\s*\(?\s*expected\s+graduation(?:\s+in|:)?\s*[0-9A-Za-z ./-]+\)?", "", text, flags=re.I)

        education_terms = re.compile(
            r"\b(?:university|college|institute|academy|school|center|centre|training|bootcamp|polytechnic|"
            r"bachelor|master|phd|doctor|degree|diploma|certificate|certification|course|major|faculty|"
            r"computer science|information technology|software engineering|graduat(?:e|ion))\b",
            re.I,
        )
        reject_terms = re.compile(
            r"\b(?:python|javascript|typescript|html|css|pytorch|cnn|transformer|ocr|rag|langchain|langgraph|"
            r"fastapi|opencv|docker|github|aws|sql|linux|developer|engineer|project|experience|skill|"
            r"english|japanese|chinese|korean|ielts|toeic|toefl|jlpt)\b",
            re.I,
        )
        named_provider = re.compile(r"\b[A-Z]{2,}(?:\s+[A-Z][A-Za-z0-9&.-]+){0,3}\b")
        parts = [
            part.strip(" ,;|()-")
            for part in re.split(r"\s*(?:\||;|\n|,)\s*", text)
            if part.strip(" ,;|()-")
        ]
        kept = []
        seen = set()
        for part in parts:
            normalized = self._clean_heading(part)
            if not normalized or normalized in seen:
                continue
            looks_like_provider = bool(named_provider.search(part)) and len(part.split()) <= 6
            if (education_terms.search(part) or looks_like_provider) and not reject_terms.search(part):
                kept.append(part)
                seen.add(normalized)

        summary = ", ".join(kept[:2])
        if summary and graduation:
            summary = f"{summary} (Expected Graduation: {graduation})"
        return summary[:240] if summary else "Not Found"

    def _coerce_float(self, value, default: float) -> float:
        try:
            return float(value)
        except (TypeError, ValueError):
            return default

    def _normalize_whitespace(self, text: str) -> str:
        text = text.replace("\r", "\n")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _split_sections(self, text: str) -> dict:
        sections = {"raw": text}
        current = "header"
        sections[current] = []

        heading_map = {}
        for section, aliases in SECTION_ALIASES.items():
            for alias in aliases:
                heading_map[self._clean_heading(alias)] = section

        for raw_line in text.splitlines():
            line = raw_line.strip(" :-|").strip()
            if not line:
                continue
            cells = self._split_table_cells(line)
            if cells and self._clean_heading(cells[0]) == "project name":
                current = "projects"
                sections.setdefault(current, []).append(line)
                continue
            cleaned = self._clean_heading(line)
            if cleaned in heading_map or self._looks_like_heading(cleaned, heading_map):
                closest = cleaned if cleaned in heading_map else self._closest_heading(cleaned, heading_map)
                current = heading_map[closest]
                sections.setdefault(current, [])
                continue
            sections.setdefault(current, []).append(line)

        return {
            key: "\n".join(value) if isinstance(value, list) else value
            for key, value in sections.items()
        }

    def _clean_heading(self, text: str) -> str:
        text = self._strip_accents(text).lower()
        text = re.sub(r"[^a-z0-9 ]+", " ", text)
        return re.sub(r"\s+", " ", text).strip()

    def _looks_like_heading(self, cleaned: str, heading_map: dict) -> bool:
        if len(cleaned.split()) > 4:
            return False
        return any(SequenceMatcher(None, cleaned, heading).ratio() >= 0.88 for heading in heading_map)

    def _closest_heading(self, cleaned: str, heading_map: dict) -> str:
        return max(heading_map, key=lambda heading: SequenceMatcher(None, cleaned, heading).ratio())

    def _extract_candidate_name(self, text: str) -> str:
        lines = [line.strip() for line in text.splitlines()[:12] if line.strip()]
        for line in lines:
            if self._is_contact_line(line) or self._is_probably_title(line):
                continue
            if self._is_table_noise_line(line):
                continue
            line = re.sub(r"\s*[-|]\s*(resume|cv|curriculum vitae)\s*$", "", line, flags=re.I)
            cleaned = re.sub(r"[^A-Za-zÀ-ỹ .'-]", "", line).strip()
            words = cleaned.split()
            if 2 <= len(words) <= 5:
                return " ".join(word.capitalize() if word.isupper() else word for word in words)
        return "Candidate"

    def _is_contact_line(self, line: str) -> bool:
        return bool(re.search(r"@|https?://|linkedin|github|phone|mobile|\+?\d[\d\s().-]{7,}", line, re.I))

    def _is_probably_title(self, line: str) -> bool:
        normalized = self._clean_heading(line)
        return any(term in normalized.split() for term in ROLE_TITLE_TERMS) and len(normalized.split()) <= 7

    def _extract_skills(self, text: str, sections: dict) -> list:
        search_text = "\n".join(
            [
                sections.get("skills", ""),
                sections.get("summary", ""),
                sections.get("experience", ""),
                sections.get("projects", ""),
                text,
            ]
        ).lower()

        found = {}
        for canonical, aliases in SKILL_TAXONOMY.items():
            for alias in aliases:
                pattern = r"(?<![a-z0-9])" + re.escape(alias.lower()) + r"(?![a-z0-9])"
                match = re.search(pattern, search_text)
                if match:
                    found[canonical] = min(found.get(canonical, match.start()), match.start())
                    break

        if "LLM APIs" in found and "LLM" in found and abs(found["LLM APIs"] - found["LLM"]) <= 2:
            found.pop("LLM", None)

        return [
            canonical
            for canonical, _position in sorted(found.items(), key=lambda item: (item[1], list(SKILL_TAXONOMY).index(item[0])))
        ]

    def _extract_recent_role(self, text: str, sections: dict) -> str:
        experience = sections.get("experience", "")
        projects = sections.get("projects", "")
        role_source = "\n".join(part for part in [experience, projects] if part.strip()) or text
        candidates = []

        project_role = self._extract_project_role(projects)
        if project_role:
            candidates.append(project_role)

        for line in role_source.splitlines()[:120]:
            candidate = self._role_from_position_row(line)
            if candidate:
                candidates.append(candidate)

        for line in role_source.splitlines()[:60]:
            if self._is_contact_line(line) or self._is_table_noise_line(line) or self._is_section_heading(line):
                continue
            for cell in self._split_table_cells(line):
                if self._looks_like_role_value(cell):
                    candidates.append(cell.strip(" -*•|"))

        if candidates:
            return self._clean_role_line(candidates[0])[:140]

        header_lines = [line.strip() for line in text.splitlines()[:10] if line.strip()]
        for line in header_lines:
            if self._is_probably_title(line) and not self._is_section_heading(line):
                return self._clean_role_line(line)[:140]
        return "Not Found"

    def _extract_project_role(self, projects: str) -> str:
        if not projects.strip():
            return ""

        project_name = ""
        position = ""
        for line in projects.splitlines()[:80]:
            cells = self._split_table_cells(line)
            if len(cells) >= 2:
                label = self._clean_heading(cells[0])
                value = cells[1].strip()
                if label == "project name" and not project_name:
                    project_name = value
                elif label in {"position", "positions", "role", "title"} and not position:
                    position = self._clean_role_line(value)
            else:
                cleaned = self._clean_heading(line)
                if not project_name and re.search(r"\b(project|crawler|bot|ocr|rag|ai|ml|vision)\b", cleaned):
                    project_name = line.strip(" -*•|")

        if not project_name:
            return ""

        project_signal = self._clean_heading(project_name)
        project_is_ai = any(term in project_signal for term in ["ai", "ml", "bot", "crawler", "ocr", "rag", "vision", "model"])
        role = position or "Developer"
        role_words = self._clean_heading(role).split()
        if project_is_ai and any(word in role_words for word in ["backend", "frontend", "fullstack", "software"]):
            role = "Developer"
        if self._looks_like_role_value(role) or role == "Developer":
            return f"{role} (Mock Project:{project_name})"[:140]
        return ""

    def _clean_role_line(self, line: str) -> str:
        return re.sub(r"^\s*(position|title|role|current role)\s*:\s*", "", line, flags=re.I).strip()

    def _role_from_position_row(self, line: str) -> str:
        cells = self._split_table_cells(line)
        if len(cells) < 2:
            return ""
        label = self._clean_heading(cells[0])
        if label not in {"position", "positions", "role", "title", "current role"}:
            return ""
        for cell in cells[1:]:
            if self._looks_like_role_value(cell):
                return cell.strip()
        return ""

    def _split_table_cells(self, line: str) -> list:
        cells = [cell.strip() for cell in line.split("|")] if "|" in line else [line.strip()]
        return self._dedupe_cells(cells)

    def _dedupe_cells(self, cells: list) -> list:
        deduped = []
        seen = set()
        for cell in cells:
            normalized = self._clean_heading(cell)
            if not normalized or normalized in seen:
                continue
            deduped.append(cell)
            seen.add(normalized)
        return deduped

    def _is_table_noise_line(self, line: str) -> bool:
        normalized = self._clean_heading(line)
        if not normalized:
            return True
        cells = [self._clean_heading(cell) for cell in self._split_table_cells(line)]
        if len(cells) >= 2 and cells[0] in TABLE_LABELS:
            return True
        label_hits = sum(1 for cell in cells if cell in TABLE_LABELS)
        return label_hits >= 2

    def _looks_like_role_value(self, text: str) -> bool:
        normalized = self._clean_heading(self._clean_role_line(text))
        words = normalized.split()
        if not words or len(words) > 8:
            return False
        if normalized in TABLE_LABELS:
            return False
        if any(term in words for term in ROLE_TITLE_TERMS):
            return True
        return any(
            re.search(rf"(?<![a-z0-9]){re.escape(self._clean_heading(title))}(?![a-z0-9])", normalized)
            for profile in ROLE_PROFILES.values()
            for title in profile["title"]
        )

    def _is_section_heading(self, line: str) -> bool:
        normalized = self._clean_heading(line)
        return any(normalized == self._clean_heading(alias) for aliases in SECTION_ALIASES.values() for alias in aliases)

    def _role_terms(self) -> set:
        terms = set()
        for profile in ROLE_PROFILES.values():
            for title in profile["title"]:
                terms.update(self._clean_heading(title).split())
        return terms | ROLE_TITLE_TERMS

    def _estimate_years_experience(self, text: str, sections: dict) -> float:
        explicit = self._extract_explicit_years(text)
        experience_text = sections.get("experience", "")
        ranges = self._extract_date_ranges(experience_text)
        months = self._merge_ranges_in_months(ranges)

        if months:
            range_years = round(months / 12, 1)
            return max(range_years, explicit or 0.0)
        if explicit:
            return explicit
        return 0.0

    def _extract_explicit_years(self, text: str) -> float:
        patterns = [
            r"(\d+(?:\.\d+)?)\+?\s*(?:years|yrs)\s+(?:of\s+)?experience",
            r"experience\s*:?\s*(\d+(?:\.\d+)?)\+?\s*(?:years|yrs)",
        ]
        values = []
        for pattern in patterns:
            for match in re.finditer(pattern, text, re.I):
                values.append(float(match.group(1)))
        return max(values) if values else 0.0

    def _extract_date_ranges(self, text: str) -> list:
        date_token = (
            r"(?:(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\.?\s+)?"
            r"(?:20\d{2}|19\d{2})"
            r"|(?:0?[1-9]|1[0-2])[/.-](?:20\d{2}|19\d{2})"
        )
        range_pattern = re.compile(
            rf"({date_token})\s*(?:-|–|—|to|until)\s*((?:present|current|now)|{date_token})",
            re.I,
        )
        ranges = []
        for match in range_pattern.finditer(text):
            context = self._date_range_context(text, match.start(), match.end())
            if self._is_education_date_context(context):
                continue
            start = self._parse_date_token(match.group(1))
            end = self._parse_date_token(match.group(2), end=True)
            if start and end and start <= end:
                ranges.append((start, end))
        return ranges

    def _date_range_context(self, text: str, start: int, end: int) -> str:
        line_start = text.rfind("\n", 0, start) + 1
        line_end = text.find("\n", end)
        if line_end == -1:
            line_end = len(text)
        return text[line_start:line_end].lower()

    def _is_education_date_context(self, context: str) -> bool:
        normalized = self._clean_heading(context)
        if not normalized:
            return False
        has_education_context = any(term in normalized for term in EDUCATION_DATE_CONTEXT_TERMS)
        has_work_context = any(term in normalized for term in WORK_DATE_CONTEXT_TERMS)
        return has_education_context and not has_work_context

    def _parse_date_token(self, token: str, end: bool = False):
        token = token.strip().lower().replace(".", "")
        if token in {"present", "current", "now"}:
            today = date.today()
            return today.year * 12 + today.month

        month_year = re.search(r"([a-z]+)\s+(19\d{2}|20\d{2})", token)
        if month_year:
            month = MONTHS.get(month_year.group(1)[:3], 1)
            return int(month_year.group(2)) * 12 + month

        numeric = re.search(r"(0?[1-9]|1[0-2])[/.-](19\d{2}|20\d{2})", token)
        if numeric:
            return int(numeric.group(2)) * 12 + int(numeric.group(1))

        year = re.search(r"(19\d{2}|20\d{2})", token)
        if year:
            return int(year.group(1)) * 12 + (12 if end else 1)
        return None

    def _merge_ranges_in_months(self, ranges: list) -> int:
        if not ranges:
            return 0
        ranges = sorted(ranges)
        merged = []
        for start, end in ranges:
            if not merged or start > merged[-1][1] + 1:
                merged.append([start, end])
            else:
                merged[-1][1] = max(merged[-1][1], end)
        return sum(end - start + 1 for start, end in merged)

    def _infer_role(self, text: str, recent_role: str, skills: list) -> tuple[str, float]:
        skill_set = set(skills)
        title_text = self._clean_heading(recent_role)
        full_text = self._clean_heading(text[:2500])
        scored = []

        for role, profile in ROLE_PROFILES.items():
            profile_skills = set(profile["skills"])
            skill_score = len(skill_set & profile_skills) / max(math.sqrt(len(profile_skills)), 1)
            title_score = 0.0
            for title in profile["title"]:
                cleaned_title = self._clean_heading(title)
                if cleaned_title in title_text:
                    title_score = max(title_score, 0.5 if cleaned_title in GENERIC_ROLE_ALIASES else 2.0)
                elif cleaned_title in full_text:
                    title_score = max(title_score, 0.25 if cleaned_title in GENERIC_ROLE_ALIASES else 1.0)
            scored.append((role, skill_score + title_score))

        scored.sort(key=lambda item: item[1], reverse=True)
        best_role, best_score = scored[0]
        second_score = scored[1][1] if len(scored) > 1 else 0.0
        confidence = min(1.0, (best_score - second_score + 1.0) / 3.0)
        return best_role, round(confidence, 2)

    def _infer_level(self, years: float, recent_role: str, text: str) -> int:
        role_text = self._clean_heading(f"{recent_role} {text[:1000]}")
        if any(term in role_text for term in ["principal", "staff", "architect", "head of", "director"]):
            return 4
        if any(term in role_text for term in ["lead", "manager", "senior"]) and years >= 3:
            return 3
        if years >= 8:
            return 4
        if years >= 5:
            return 3
        if years >= 2:
            return 2
        return 1

    def _extract_education(self, sections: dict) -> str:
        education = sections.get("education", "").strip()
        if not education:
            return "Not Found"
        lines = [line.strip(" -*•") for line in education.splitlines() if line.strip()]
        return " | ".join(lines[:3])[:240] if lines else "Not Found"

    def _strip_accents(self, text: str) -> str:
        import unicodedata

        stripped = "".join(
            char for char in unicodedata.normalize("NFKD", text) if not unicodedata.combining(char)
        )
        return stripped.replace("đ", "d").replace("Đ", "D")


def _extract_education_v2(self, sections: dict) -> str:
    education = sections.get("education", "").strip()
    if not education:
        return "Not Found"
    cells = []
    expected_graduation = ""
    for line in education.splitlines():
        for cell in self._split_table_cells(line):
            cleaned = self._clean_education_cell(cell)
            if not cleaned:
                continue
            grad = self._extract_expected_graduation(cleaned)
            if grad and not expected_graduation:
                expected_graduation = grad
                continue
            cells.append(cleaned)

    unique_cells = []
    seen = set()
    for cell in cells:
        normalized = self._clean_heading(cell)
        if normalized and normalized not in seen:
            unique_cells.append(cell)
            seen.add(normalized)

    if not unique_cells and not expected_graduation:
        return "Not Found"

    summary = ", ".join(unique_cells[:2])
    if expected_graduation:
        summary = f"{summary} (Expected Graduation: {expected_graduation})" if summary else f"Expected Graduation: {expected_graduation}"
    return summary[:240]


def _clean_education_cell(self, cell: str) -> str:
    value = cell.strip(" -*â€¢|")
    if not value:
        return ""
    normalized = self._clean_heading(value)
    if normalized in {"education", "academic background", "qualification", "qualifications"}:
        return ""
    if re.search(r"\b(?:english|japanese|chinese|korean|french|german|spanish|vietnamese)\s*:?\s*[a-z0-9.+-]+(?:\s+level)?\b", value, re.I):
        return ""
    if re.search(r"\bfrom\s+\d{1,2}[/.-]\d{4}\s+to\s*$", value, re.I):
        return ""
    if re.search(r"\b(?:gpa|ielts|toeic|toefl)\s*:?", value, re.I):
        return ""
    value = re.sub(r"\s*\bfrom\s+\d{1,2}[/.-]\d{4}\s+to\s*$", "", value, flags=re.I)
    return re.sub(r"\s+", " ", value).strip(" ,")


def _extract_expected_graduation(self, text: str) -> str:
    match = re.search(
        r"expected\s+graduation(?:\s+in|:)?\s*([0-9]{1,2}[/.-][0-9]{4}|[A-Za-z]+\s+[0-9]{4}|[0-9]{4})",
        text,
        re.I,
    )
    return match.group(1).strip() if match else ""


CVExtractor._extract_education = _extract_education_v2
CVExtractor._clean_education_cell = _clean_education_cell
CVExtractor._extract_expected_graduation = _extract_expected_graduation

cv_extractor = CVExtractor()

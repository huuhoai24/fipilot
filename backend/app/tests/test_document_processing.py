from __future__ import annotations

import io
import time
import zipfile

import pytest
from docx import Document
from pypdf import PdfWriter

from infrastructure.documents import (
    DocumentExtractionStatus,
    DocumentProcessingError,
    DocumentService,
)


def _docx_bytes(text: str = "Candidate\nEXPERIENCE\nEngineer at Example Co") -> bytes:
    buffer = io.BytesIO()
    document = Document()
    document.add_paragraph(text)
    document.save(buffer)
    return buffer.getvalue()


def test_rejects_extension_and_content_mismatch(tmp_path):
    path = tmp_path / "resume.pdf"
    path.write_bytes(_docx_bytes())

    with pytest.raises(DocumentProcessingError) as caught:
        DocumentService().extract_document(
            str(path), "resume.pdf", content_type="application/pdf"
        )

    assert caught.value.code == "file_type_mismatch"
    assert caught.value.status_code == 415


def test_rejects_malformed_docx_container(tmp_path):
    path = tmp_path / "resume.docx"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("unrelated.txt", "not a Word document")

    with pytest.raises(DocumentProcessingError) as caught:
        DocumentService().extract_document(str(path), "resume.docx")

    assert caught.value.code == "invalid_document"
    assert caught.value.status_code == 422


def test_docx_result_separates_document_metadata_and_tables(tmp_path):
    path = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("Candidate Name - Backend Engineer with production delivery experience")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "FastAPI"
    document.save(path)

    result = DocumentService().extract_document(str(path), "resume.docx")

    assert result.status is DocumentExtractionStatus.COMPLETE
    assert result.source_type == "docx"
    assert result.character_count == len(result.text)
    assert result.page_count is None
    assert result.tables[0].rows == [["Python", "FastAPI"]]
    assert "Python | FastAPI" in result.text


class _FakeOCR:
    def recognize(self, image_bytes: bytes) -> str:
        assert image_bytes
        return "Scanned Candidate\nEXPERIENCE\nEngineer at OCR Systems | 2022 - Present"


def test_image_only_pdf_uses_injected_ocr_and_marks_method(tmp_path):
    path = tmp_path / "scan.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=300, height=300)
    with path.open("wb") as stream:
        writer.write(stream)

    result = DocumentService(ocr_engine=_FakeOCR()).extract_document(
        str(path), "scan.pdf"
    )

    assert result.status is DocumentExtractionStatus.COMPLETE
    assert result.extraction_method == "ocr"
    assert result.pages[0].extraction_method == "ocr"
    assert "Scanned Candidate" in result.text


class _FailingOCR:
    def recognize(self, image_bytes: bytes) -> str:
        raise RuntimeError("private OCR failure")


def test_ocr_failure_is_structured_when_no_text_can_be_recovered(tmp_path):
    path = tmp_path / "scan.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=300, height=300)
    with path.open("wb") as stream:
        writer.write(stream)

    with pytest.raises(DocumentProcessingError) as caught:
        DocumentService(ocr_engine=_FailingOCR()).extract_document(str(path), "scan.pdf")

    assert caught.value.code == "no_extractable_text"
    assert "ocr_failed" in caught.value.warnings
    assert "private" not in caught.value.safe_message.lower()


class _SlowOCR:
    def recognize(self, image_bytes: bytes) -> str:
        time.sleep(0.05)
        return "Text returned too late to be accepted by the bounded OCR stage."


def test_ocr_timeout_has_a_distinct_safe_warning(tmp_path):
    path = tmp_path / "scan.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=300, height=300)
    with path.open("wb") as stream:
        writer.write(stream)

    with pytest.raises(DocumentProcessingError) as caught:
        DocumentService(ocr_engine=_SlowOCR(), ocr_timeout_seconds=0.001).extract_document(
            str(path), "scan.pdf"
        )

    assert caught.value.code == "no_extractable_text"
    assert "ocr_timeout" in caught.value.warnings

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPORT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = REPORT_ROOT / "Final_report.docx"
FIGURE_ROOT = REPORT_ROOT / "figures"

CHAPTER_SOURCES = [
    REPORT_ROOT / "chapters" / "01_project_introduction.tex",
    REPORT_ROOT / "chapters" / "02_project_management.tex",
    REPORT_ROOT / "chapters" / "03_theoretical_background.tex",
    REPORT_ROOT / "chapters" / "04_ai_methodology.tex",
    REPORT_ROOT / "chapters" / "05_system_design.tex",
    REPORT_ROOT / "chapters" / "06_requirements.tex",
    REPORT_ROOT / "chapters" / "07_software_testing.tex",
    REPORT_ROOT / "chapters" / "08_deliverables_user_guide.tex",
    REPORT_ROOT / "chapters" / "09_conclusion.tex",
]

CITATIONS = {
    "vaswani2017attention": "[1]",
    "lewis2020rag": "[2]",
    "zheng2023judge": "[3]",
}

REFERENCES = [
    "[1] A. Vaswani et al., “Attention Is All You Need,” Advances in Neural Information Processing Systems, vol. 30, 2017.",
    "[2] P. Lewis et al., “Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks,” Advances in Neural Information Processing Systems, vol. 33, 2020.",
    "[3] L. Zheng et al., “Judging LLM-as-a-Judge with MT-Bench and Chatbot Arena,” Advances in Neural Information Processing Systems, vol. 36, 2023.",
]

ROMAN = ("I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X")


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 70, start: int = 90, bottom: int = 70, end: int = 90) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep_row_together(row) -> None:
    properties = row._tr.get_or_add_trPr()
    properties.append(OxmlElement("w:cantSplit"))


def repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def add_field(paragraph, instruction: str, placeholder: str = "Update this field in Word") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set(qn("xml:space"), "preserve")
    instruction_node.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    run.extend((begin, instruction_node, separate, text, end))


def set_repeat_update_fields(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.right_margin = Inches(0.80)
    section.bottom_margin = Inches(0.80)
    section.left_margin = Inches(0.90)
    section.header_distance = Inches(0.30)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Inches(0.28)

    for name, size, color in (
        ("Title", 22, "000000"),
        ("Subtitle", 14, "244C66"),
        ("Heading 1", 16, "000000"),
        ("Heading 2", 14, "000000"),
        ("Heading 3", 12, "000000"),
    ):
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.first_line_indent = None
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].paragraph_format.page_break_before = True
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.bold = True

    caption = styles["Caption"]
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(10)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.first_line_indent = None

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code_style.paragraph_format.left_indent = Inches(0.20)
    code_style.paragraph_format.right_indent = Inches(0.20)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(6)
    code_style.paragraph_format.first_line_indent = None

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.8))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(3.4)
    table.columns[1].width = Inches(3.4)
    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left.add_run("FiPilot Project Report")
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.add_run("Editable DOCX Edition")
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
        properties = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:color"), "000000")
        borders.append(bottom)
        properties.append(borders)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    add_field(paragraph, "PAGE", "1")

    set_repeat_update_fields(document)
    document.core_properties.title = "FiPilot AI-Powered CV-to-Interview System"
    document.core_properties.subject = "Graduation Project Report"
    document.core_properties.author = "Software Engineering and Artificial Intelligence Project"


def clean_inline(text: str) -> str:
    text = text.strip()
    text = re.sub(r"\\label\{[^{}]*\}", "", text)
    text = re.sub(
        r"\\cite\{([^{}]+)\}",
        lambda match: " ".join(CITATIONS.get(key.strip(), "") for key in match.group(1).split(",")),
        text,
    )
    text = re.sub(
        r"\\cref\{([^{}]+)\}",
        lambda match: "the referenced figure" if match.group(1).startswith("fig:") else "the referenced table",
        text,
    )
    for command in ("codepath", "textbf", "textit", "emph", "text", "operatorname"):
        pattern = re.compile(rf"\\{command}\{{([^{{}}]*)\}}")
        while pattern.search(text):
            text = pattern.sub(lambda match: match.group(1), text)
    replacements = {
        r"\_": "_",
        r"\&": "&",
        r"\%": "%",
        r"\#": "#",
        r"\$": "$",
        r"\rightarrow": "→",
        r"\top": "ᵀ",
        r"\sqrt": "√",
        "~": " ",
        "---": "—",
        "--": "–",
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    text = re.sub(r"\\[A-Za-z@]+\*?(?:\[[^\]]*\])?", "", text)
    text = text.replace("{", "").replace("}", "")
    text = text.replace("$", "")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def add_body_paragraph(document: Document, text: str) -> None:
    cleaned = clean_inline(text)
    if not cleaned:
        return
    paragraph = document.add_paragraph(cleaned)
    paragraph.paragraph_format.keep_together = False


def extract_caption(block: str) -> str:
    match = re.search(r"\\caption\{([^{}]+)\}", block)
    return clean_inline(match.group(1)) if match else ""


def parse_table_rows(block: str) -> list[list[str]]:
    rows: list[list[str]] = []
    seen: set[tuple[str, ...]] = set()
    for raw_line in block.splitlines():
        line = raw_line.strip()
        if not line or line.startswith("\\begin") or line.startswith("\\end"):
            continue
        line = re.sub(r"\\caption\{[^{}]*\}", "", line)
        line = re.sub(r"\\label\{[^{}]*\}", "", line)
        for command in (r"\toprule", r"\midrule", r"\bottomrule", r"\endfirsthead", r"\endhead"):
            line = line.replace(command, "")
        line = line.strip()
        if "&" not in line:
            continue
        line = re.sub(r"\\\\\s*$", "", line).strip()
        cells = [clean_inline(cell) for cell in re.split(r"(?<!\\)&", line)]
        cells = [cell for cell in cells if cell or len(cells) > 1]
        key = tuple(cells)
        if len(cells) < 2 or key in seen:
            continue
        seen.add(key)
        rows.append(cells)
    return rows


def add_table(document: Document, block: str, table_number: str) -> None:
    caption = extract_caption(block)
    if caption:
        paragraph = document.add_paragraph(style="Caption")
        run = paragraph.add_run(f"Table {table_number}: ")
        run.bold = True
        paragraph.add_run(caption)
    rows = parse_table_rows(block)
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, source_row in enumerate(rows):
        target_row = table.rows[row_index]
        keep_row_together(target_row)
        if row_index == 0:
            repeat_table_header(target_row)
        for column_index in range(column_count):
            cell = target_row.cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            value = source_row[column_index] if column_index < len(source_row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            if row_index == 0:
                run.bold = True
                set_cell_shading(cell, "E9EFF3")
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(document: Document, filename: str, caption: str, number: str) -> None:
    path = FIGURE_ROOT / filename
    if not path.exists():
        raise FileNotFoundError(f"Missing report figure: {path}")
    from docx.image.image import Image

    image = Image.from_file(str(path))
    ratio = image.px_width / image.px_height
    max_width = 6.25
    max_height = 6.80
    if max_width / ratio <= max_height:
        width = max_width
        height = max_width / ratio
    else:
        height = max_height
        width = max_height * ratio
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width), height=Inches(height))
    caption_paragraph = document.add_paragraph(style="Caption")
    caption_run = caption_paragraph.add_run(f"Figure {number}: ")
    caption_run.bold = True
    caption_paragraph.add_run(clean_inline(caption))
    document.add_page_break()


def normalize_equation(raw: str) -> str:
    if "Resume" in raw and "Profile" in raw:
        return "Resume → Profile → Plan → Question → Answer → Evaluation → Report"
    if "Attention" in raw:
        return "Attention(Q, K, V) = softmax(QKᵀ / √dₖ) V"
    return clean_inline(raw)


class TexExporter:
    def __init__(self, document: Document, appendix: bool = False) -> None:
        self.document = document
        self.appendix = appendix
        self.chapter = 0
        self.section = 0
        self.subsection = 0
        self.figure = 0
        self.table = 0

    @property
    def chapter_label(self) -> str:
        if self.appendix:
            return chr(ord("A") + self.chapter - 1)
        return ROMAN[self.chapter - 1]

    def heading(self, level: int, title: str, numbered: bool = True) -> None:
        title = clean_inline(title)
        if level == 1:
            self.chapter += 1
            self.section = self.subsection = self.figure = self.table = 0
            prefix = f"APPENDIX {self.chapter_label}." if self.appendix else f"CHAPTER {self.chapter_label}."
            text = f"{prefix} {title}" if numbered else title
        elif level == 2:
            self.section += 1
            self.subsection = 0
            text = f"{self.chapter_label}.{self.section} {title}" if numbered else title
        else:
            self.subsection += 1
            text = f"{self.chapter_label}.{self.section}.{self.subsection} {title}" if numbered else title
        self.document.add_heading(text, level=level)

    def next_figure_number(self) -> str:
        self.figure += 1
        return f"{self.chapter_label}.{self.figure}"

    def next_table_number(self) -> str:
        self.table += 1
        return f"{self.chapter_label}.{self.table}"

    def export(self, path: Path, frontmatter: bool = False) -> None:
        lines = path.read_text(encoding="utf-8").splitlines()
        paragraph_lines: list[str] = []

        def flush_paragraph() -> None:
            if paragraph_lines:
                add_body_paragraph(self.document, " ".join(paragraph_lines))
                paragraph_lines.clear()

        index = 0
        while index < len(lines):
            raw = lines[index]
            line = raw.strip()
            if not line or line.startswith("%"):
                flush_paragraph()
                index += 1
                continue

            heading_match = re.match(r"\\(chapter|section|subsection)(\*)?\{([^{}]+)\}", line)
            if heading_match:
                flush_paragraph()
                kind, starred, title = heading_match.groups()
                level = {"chapter": 1, "section": 2, "subsection": 3}[kind]
                if frontmatter:
                    self.document.add_heading(clean_inline(title), level=1)
                else:
                    self.heading(level, title, numbered=not bool(starred))
                index += 1
                continue

            if line.startswith("\\addcontentsline") or line.startswith("\\label"):
                index += 1
                continue

            if line.startswith("\\reportfigure"):
                flush_paragraph()
                match = re.search(r"\\reportfigure(?:\[[^\]]+\])?\{([^{}]+)\}\{([^{}]+)\}\{([^{}]+)\}", line)
                if not match:
                    raise ValueError(f"Cannot parse figure line in {path}: {line}")
                filename, caption, _label = match.groups()
                add_figure(self.document, filename, caption, self.next_figure_number())
                index += 1
                continue

            if line.startswith("\\begin{table") or line.startswith("\\begin{longtable"):
                flush_paragraph()
                end_token = "\\end{table}" if line.startswith("\\begin{table") else "\\end{longtable}"
                block = [raw]
                index += 1
                while index < len(lines):
                    block.append(lines[index])
                    if end_token in lines[index]:
                        break
                    index += 1
                add_table(self.document, "\n".join(block), self.next_table_number())
                index += 1
                continue

            if line.startswith("\\begin{lstlisting}"):
                flush_paragraph()
                code_lines: list[str] = []
                caption_match = re.search(r"caption=\{([^{}]+)\}", line)
                if caption_match:
                    caption = self.document.add_paragraph(style="Caption")
                    caption.add_run(clean_inline(caption_match.group(1))).bold = True
                index += 1
                while index < len(lines) and not lines[index].strip().startswith("\\end{lstlisting}"):
                    code_lines.append(lines[index])
                    index += 1
                paragraph = self.document.add_paragraph(style="Code Block")
                paragraph.add_run("\n".join(code_lines))
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "F2F5F7")
                paragraph._p.get_or_add_pPr().append(shading)
                index += 1
                continue

            if line.startswith("\\begin{equation}"):
                flush_paragraph()
                equation_lines: list[str] = []
                index += 1
                while index < len(lines) and not lines[index].strip().startswith("\\end{equation}"):
                    equation_lines.append(lines[index])
                    index += 1
                paragraph = self.document.add_paragraph(normalize_equation(" ".join(equation_lines)))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = None
                index += 1
                continue

            if line.startswith("\\begin{itemize}") or line.startswith("\\begin{enumerate}"):
                flush_paragraph()
                ordered = "enumerate" in line
                end_token = "\\end{enumerate}" if ordered else "\\end{itemize}"
                index += 1
                item_lines: list[str] = []
                while index < len(lines) and not lines[index].strip().startswith(end_token):
                    item_line = lines[index].strip()
                    if item_line.startswith("\\item"):
                        if item_lines:
                            paragraph = self.document.add_paragraph(clean_inline(" ".join(item_lines)), style="List Number" if ordered else "List Bullet")
                            paragraph.paragraph_format.first_line_indent = None
                        item_lines = [item_line[len("\\item") :].strip()]
                    elif item_line:
                        item_lines.append(item_line)
                    index += 1
                if item_lines:
                    paragraph = self.document.add_paragraph(clean_inline(" ".join(item_lines)), style="List Number" if ordered else "List Bullet")
                    paragraph.paragraph_format.first_line_indent = None
                index += 1
                continue

            if line.startswith("\\begin") or line.startswith("\\end"):
                index += 1
                continue

            paragraph_lines.append(line)
            index += 1

        flush_paragraph()


def collect_lists(paths: Iterable[Path]) -> tuple[list[str], list[str]]:
    figures: list[str] = []
    tables: list[str] = []
    for path in paths:
        text = path.read_text(encoding="utf-8")
        for match in re.finditer(r"\\reportfigure(?:\[[^\]]+\])?\{[^{}]+\}\{([^{}]+)\}", text):
            figures.append(clean_inline(match.group(1)))
        for match in re.finditer(r"\\caption\{([^{}]+)\}", text):
            tables.append(clean_inline(match.group(1)))
    return figures, tables


def add_title_page(document: Document) -> None:
    for _ in range(4):
        document.add_paragraph()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    run = paragraph.add_run("FIPILOT")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(26)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    run = paragraph.add_run("AI-POWERED CV-TO-INTERVIEW SYSTEM")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(18)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    run = paragraph.add_run("GRADUATION PROJECT REPORT")
    run.font.name = "Times New Roman"
    run.font.size = Pt(15)
    for _ in range(8):
        document.add_paragraph()
    paragraph = document.add_paragraph("Software Engineering and Artificial Intelligence Project")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph = document.add_paragraph("Quy Nhon, August 2026")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    document.add_page_break()


def add_front_lists(document: Document, figures: list[str], tables: list[str]) -> None:
    document.add_heading("Table of Contents", level=1)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None
    add_field(paragraph, r'TOC \o "1-3" \h \z \u', "Right-click and choose Update Field in Word")

    document.add_heading("List of Figures", level=1)
    for index, caption in enumerate(figures, start=1):
        paragraph = document.add_paragraph(f"Figure {index}. {caption}")
        paragraph.paragraph_format.first_line_indent = None

    document.add_heading("List of Tables", level=1)
    for index, caption in enumerate(tables, start=1):
        paragraph = document.add_paragraph(f"Table {index}. {caption}")
        paragraph.paragraph_format.first_line_indent = None


def build() -> Path:
    document = Document()
    configure_document(document)
    add_title_page(document)

    front_exporter = TexExporter(document)
    front_exporter.export(REPORT_ROOT / "frontmatter.tex", frontmatter=True)

    all_content_paths = [*CHAPTER_SOURCES, REPORT_ROOT / "appendices" / "runtime_appendices.tex"]
    figures, tables = collect_lists(all_content_paths)
    add_front_lists(document, figures, tables)

    exporter = TexExporter(document)
    for source in CHAPTER_SOURCES:
        exporter.export(source)

    document.add_heading("REFERENCES", level=1)
    for reference in REFERENCES:
        paragraph = document.add_paragraph(reference)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.left_indent = Inches(0.25)

    appendix_exporter = TexExporter(document, appendix=True)
    appendix_exporter.export(REPORT_ROOT / "appendices" / "runtime_appendices.tex")

    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    result = build()
    print(result)
